import logging
from io import BytesIO
from datetime import datetime
import sentry_sdk
from docx import Document
from reportlab.lib.pagesizes import LETTER
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT
from typing import Dict, Any

logger = logging.getLogger(__name__)

class ExportService:
    def generate_docx(self, proposal: Dict[str, Any]) -> BytesIO:
        """
        Generates a professional DOCX file from proposal section data.
        """
        timestamp = datetime.utcnow().isoformat()
        logger.info(f"timestamp={timestamp}, service=ExportService, operation=generate_docx, status=started")
        
        try:
            document = Document()
            
            # Extract sections
            sections = proposal.get("sections", [])
            
            # Add sections to the document
            for section in sections:
                title = section.get("title", "Untitled Section")
                content = section.get("content", "")
                
                # Add Header
                document.add_heading(title, level=1)
                
                # Add Content
                document.add_paragraph(content)
                
                # Add spacing between sections
                document.add_paragraph("")

            # Save to BytesIO buffer
            buffer = BytesIO()
            document.save(buffer)
            buffer.seek(0)
            
            logger.info(f"timestamp={timestamp}, service=ExportService, operation=generate_docx, status=success")
            return buffer
            
        except Exception as e:
            sentry_sdk.capture_exception(e)
            logger.error(f"timestamp={timestamp}, service=ExportService, operation=generate_docx, status=error, error={str(e)}")
            raise

    def generate_pdf(self, proposal: Dict[str, Any]) -> BytesIO:
        """
        Generates a professional PDF file from proposal section data.
        """
        timestamp = datetime.utcnow().isoformat()
        logger.info(f"timestamp={timestamp}, service=ExportService, operation=generate_pdf, status=started")
        
        try:
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=LETTER)
            styles = getSampleStyleSheet()
            
            # Custom style for section titles
            title_style = ParagraphStyle(
                'SectionTitle',
                parent=styles['Heading1'],
                fontSize=14,
                spaceAfter=12,
                bold=True
            )
            
            # Custom style for body text
            body_style = styles['Normal']
            body_style.alignment = TA_LEFT
            body_style.spaceAfter = 12
            
            elements = []
            
            # Extract sections
            sections = proposal.get("sections", [])
            
            for section in sections:
                title = section.get("title", "Untitled Section")
                content = section.get("content", "")
                
                # Add Heading
                elements.append(Paragraph(title, title_style))
                
                # Add Paragraphs (handle newlines)
                for part in content.split('\n'):
                    if part.strip():
                        elements.append(Paragraph(part, body_style))
                
                # Space after section
                elements.append(Spacer(1, 12))

            # Build PDF
            doc.build(elements)
            buffer.seek(0)
            
            logger.info(f"timestamp={timestamp}, service=ExportService, operation=generate_pdf, status=success")
            return buffer
            
        except Exception as e:
            sentry_sdk.capture_exception(e)
            logger.error(f"timestamp={timestamp}, service=ExportService, operation=generate_pdf, status=error, error={str(e)}")
            raise
